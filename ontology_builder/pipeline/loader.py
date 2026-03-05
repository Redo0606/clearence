"""Load document text from PDF, DOCX, TXT, or MD. Raises ValueError for unsupported format."""

import logging
from pathlib import Path

import docx
from pdfminer.high_level import extract_text as pdfminer_extract_text

logger = logging.getLogger(__name__)


def load_document(path: str) -> str:
    """Load document text from PDF, DOCX, TXT, or MD.

    Args:
        path: File path to the document.

    Returns:
        Extracted text content.

    Raises:
        FileNotFoundError: If path does not exist.
        ValueError: If file format is unsupported.
        RuntimeError: If extraction fails.
    """
    p = Path(path)
    logger.debug("[Loader] Loading document | path=%s", path)
    if not p.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    suffix = p.suffix.lower()
    logger.debug("[Loader] Detected format | suffix=%s", suffix)

    if suffix == ".pdf":
        try:
            logger.debug("[Loader] Extracting text via pdfminer")
            text = pdfminer_extract_text(path)
            logger.info("[Loader] PDF loaded | chars=%d", len(text))
            return text
        except Exception as e:
            raise RuntimeError(f"Failed to extract PDF text: {e}") from e

    if suffix == ".docx":
        try:
            logger.debug("[Loader] Reading DOCX paragraphs")
            doc = docx.Document(path)
            text = "\n".join(para.text for para in doc.paragraphs)
            logger.info("[Loader] DOCX loaded | chars=%d | paragraphs=%d", len(text), len(doc.paragraphs))
            return text
        except Exception as e:
            raise RuntimeError(f"Failed to read DOCX: {e}") from e

    if suffix in (".txt", ".md"):
        try:
            logger.debug("[Loader] Reading plain text")
            text = p.read_text(encoding="utf-8", errors="replace")
            logger.info("[Loader] Text file loaded | chars=%d", len(text))
            return text
        except Exception as e:
            raise RuntimeError(f"Failed to read file: {e}") from e

    raise ValueError(f"Unsupported format: {suffix}. Use .pdf, .docx, .txt, or .md")
